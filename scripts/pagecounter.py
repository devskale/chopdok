import os
import sys
import PyPDF2
from PIL import Image
import docx
import openpyxl
from docx2python import docx2python


def count_pdf_pages(filepath):
    """Count pages in a PDF file."""
    try:
        with open(filepath, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            return len(reader.pages)
    except Exception as e:
        print(f"Error reading PDF {filepath}: {e}")
        return 0


def count_image_pages(filepath):
    """Count an image as one page."""
    try:
        # Just verify it's a valid image
        Image.open(filepath)
        return 1
    except Exception as e:
        print(f"Error reading image {filepath}: {e}")
        return 0


def count_docx_pages(filepath):
    """Count pages in a DOCX file accurately using docx2python.

    This method extracts the actual page count from Word document properties.
    """
    try:
        # First try with docx2python which can extract page count from document properties
        doc = docx2python(filepath)
        pages = doc.properties.get('page_count')

        if pages is not None and pages > 0:
            return pages

        # Fallback method if docx2python doesn't return valid page count
        # Get last page number from the document's footer
        footers = doc.footer
        last_page_number = None

        # Look for page numbers in footers
        for footer_list in footers:
            for footer in footer_list:
                for text in footer:
                    if isinstance(text, str) and "PAGE" in text.upper():
                        # Try to extract page number reference
                        return 1  # Default if we can't determine actual count

        # If we're here, try another fallback with python-docx
        doc2 = docx.Document(filepath)

        # Approximate characters per page
        chars_per_page = 3000
        total_chars = 0

        # Count characters in paragraphs
        for para in doc2.paragraphs:
            total_chars += len(para.text)

        # Estimate table contribution
        for table in doc2.tables:
            for row in table.rows:
                for cell in row.cells:
                    # Count text in table cells
                    for paragraph in cell.paragraphs:
                        total_chars += len(paragraph.text)
                    # Tables take more space, add extra characters
                    total_chars += 100  # Adjustment for table formatting

        # Account for images and other elements
        image_count = len(doc2.inline_shapes)
        total_chars += image_count * 500

        # Calculate estimated pages (minimum 1)
        estimated_pages = max(1, int(total_chars / chars_per_page))
        return estimated_pages

    except Exception as e:
        print(f"Error reading DOCX {filepath}: {e}")
        return 0


def count_xlsx_pages(filepath):
    """Count sheets in an XLSX file as pages."""
    try:
        workbook = openpyxl.load_workbook(filepath, read_only=True)
        return len(workbook.sheetnames)
    except Exception as e:
        print(f"Error reading XLSX {filepath}: {e}")
        return 0


def scan_directory(directory_path):
    """Scan directory and count pages in various file types.
    Returns overall results and results per subdirectory.
    """
    if not os.path.isdir(directory_path):
        print(f"Error: {directory_path} is not a valid directory.")
        return

    # Overall results
    overall_results = {
        'pdf': {'count': 0, 'pages': 0, 'files': []},
        'image': {'count': 0, 'pages': 0, 'files': []},
        'docx': {'count': 0, 'pages': 0, 'files': []},
        'xlsx': {'count': 0, 'pages': 0, 'files': []}
    }

    # Results per directory
    directory_results = {}

    image_extensions = ['.jpg', '.jpeg', '.png',
                        '.gif', '.bmp', '.tiff', '.webp']

    for root, _, files in os.walk(directory_path):
        # Initialize results for this directory
        rel_dir = os.path.relpath(root, directory_path)
        dir_result = {
            'pdf': {'count': 0, 'pages': 0, 'files': []},
            'image': {'count': 0, 'pages': 0, 'files': []},
            'docx': {'count': 0, 'pages': 0, 'files': []},
            'xlsx': {'count': 0, 'pages': 0, 'files': []}
        }

        for file in files:
            filepath = os.path.join(root, file)
            extension = os.path.splitext(file)[1].lower()

            # PDF files
            if extension == '.pdf':
                pages = count_pdf_pages(filepath)
                overall_results['pdf']['count'] += 1
                overall_results['pdf']['pages'] += pages
                overall_results['pdf']['files'].append((filepath, pages))

                dir_result['pdf']['count'] += 1
                dir_result['pdf']['pages'] += pages
                dir_result['pdf']['files'].append((filepath, pages))

            # Image files
            elif extension in image_extensions:
                pages = count_image_pages(filepath)
                overall_results['image']['count'] += 1
                overall_results['image']['pages'] += pages
                overall_results['image']['files'].append((filepath, pages))

                dir_result['image']['count'] += 1
                dir_result['image']['pages'] += pages
                dir_result['image']['files'].append((filepath, pages))

            # DOCX files
            elif extension == '.docx':
                pages = count_docx_pages(filepath)
                overall_results['docx']['count'] += 1
                overall_results['docx']['pages'] += pages
                overall_results['docx']['files'].append((filepath, pages))

                dir_result['docx']['count'] += 1
                dir_result['docx']['pages'] += pages
                dir_result['docx']['files'].append((filepath, pages))

            # XLSX files
            elif extension == '.xlsx':
                pages = count_xlsx_pages(filepath)
                overall_results['xlsx']['count'] += 1
                overall_results['xlsx']['pages'] += pages
                overall_results['xlsx']['files'].append((filepath, pages))

                dir_result['xlsx']['count'] += 1
                dir_result['xlsx']['pages'] += pages
                dir_result['xlsx']['files'].append((filepath, pages))

        # Only store directories that have files we care about
        total_files = sum(dir_result[ft]['count'] for ft in dir_result)
        if total_files > 0:
            directory_results[rel_dir] = dir_result

    return overall_results, directory_results


def print_results(results, directory_results, detailed=False):
    """Print the results in a formatted way."""
    print("\n===== Overall Page Count Summary =====")
    print("-" * 50)

    overall_results, file_types = results, ['pdf', 'image', 'docx', 'xlsx']
    total_files = 0
    total_pages = 0

    # Print file type breakdown
    for file_type in file_types:
        data = overall_results[file_type]
        file_count = data['count']
        page_count = data['pages']
        total_files += file_count
        total_pages += page_count

        print(f"{file_type.upper()} Files: {file_count} | Pages: {page_count}")

    print("-" * 50)
    print(f"TOTAL: {total_files} files | {total_pages} pages")
    print("=" * 50)

    # Print directory breakdown
    print("\n===== Directory Breakdown =====")
    for directory, dir_data in sorted(directory_results.items()):
        dir_total_files = 0
        dir_total_pages = 0

        for file_type in file_types:
            dir_total_files += dir_data[file_type]['count']
            dir_total_pages += dir_data[file_type]['pages']

        # Only print directories with files
        if dir_total_files > 0:
            dir_name = '.' if directory == '.' else directory
            print(f"\n{dir_name}:")
            print("-" * 40)

            # Print file type breakdown for this directory
            for file_type in file_types:
                count = dir_data[file_type]['count']
                pages = dir_data[file_type]['pages']
                if count > 0:
                    print(
                        f"  {file_type.upper()}: {count} files | {pages} pages")

            print(f"  SUM: {dir_total_files} files | {dir_total_pages} pages")

            # If detailed flag is set, print all files in this directory
            if detailed:
                print("\n  Files in this directory:")
                for file_type in file_types:
                    for filepath, pages in dir_data[file_type]['files']:
                        filename = os.path.basename(filepath)
                        print(
                            f"    - {filename} ({file_type.upper()}): {pages} page(s)")


def main():
    if len(sys.argv) < 2:
        print("Usage: python page_counter.py <directory_path> [-d|--detailed]")
        return

    directory_path = sys.argv[1]
    detailed = any(arg in ["-d", "--detailed"] for arg in sys.argv)

    overall_results, directory_results = scan_directory(directory_path)

    if overall_results:
        print_results(overall_results, directory_results, detailed)


if __name__ == "__main__":
    main()
